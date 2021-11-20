import openpyxl
import docx

end = '\033[0m'
underline = '\033[4m'


wb = openpyxl.load_workbook('workfile.xlsx')
sheet = wb["Sheet3"]

max_col = sheet.max_row
sheet.delete_rows(sheet.min_row, 1)


for x in range(1,max_col):
    Name=sheet.cell(row=x, column=1).value
    Contact=sheet.cell(row=x, column=2).value
    Product=sheet.cell(row=x, column=4).value
    Quantity=sheet.cell(row=x, column=5).value
    Amount=sheet.cell(row=x, column=6).value+sheet.cell(row=1, column=8).value
    Date = sheet.cell(row=x, column=9).value
    print(Name,Contact,Product,Quantity,Amount)

doc= docx.Document()
doc.add_heading(Date.strftime("%b-%d"),1)


menuTable = doc.add_table(rows=1, cols=6)
menuTable.style = 'Table Grid'
hdr_Cells = menuTable.rows[0].cells
hdr_Cells[0].text = 'Name'
hdr_Cells[1].text = 'Contact'
hdr_Cells[2].text = 'Product'
hdr_Cells[3].text = 'Quantity'
hdr_Cells[4].text = 'Amount'
hdr_Cells[5].text = 'Payment'

for x in range(1, max_col):

    Name = sheet.cell(row=x, column=1).value
    Contact = sheet.cell(row=x, column=2).value
    Product = sheet.cell(row=x, column=4).value
    Quantity = sheet.cell(row=x, column=5).value
    Amount = sheet.cell(row=x, column=6).value + sheet.cell(row=x, column=8).value-sheet.cell(row=x, column=12).value
    if x>1 and Name == str(sheet.cell(row=x-1, column=1).value):
        continue
    for m in range(x+1,max_col):
        if Name == str(sheet.cell(row=m, column=1).value):
            Amount= Amount+ sheet.cell(row=m, column=6).value + sheet.cell(row=m, column=8).value
            Product= str(Product)+str(f',')+str(sheet.cell(row=m, column=4).value)
            Quantity= str(Quantity)+str(f',')+str(sheet.cell(row=m, column=5).value)

    row_Cells = menuTable.add_row().cells
    row_Cells[0].text=  str(Name)
    row_Cells[1].text = str(Contact)
    row_Cells[2].text = str(Product)
    row_Cells[3].text = str(Quantity)
    row_Cells[4].text = str(Amount)



doc.save(f'invoice\delilist\{Date.strftime("%b-%d")}.docx')



